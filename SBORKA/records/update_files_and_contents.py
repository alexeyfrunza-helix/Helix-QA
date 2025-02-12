import os
import shutil
import re
import logging
import sys
from datetime import datetime
from docx import Document

sys.stdout.reconfigure(encoding='utf-8')
# Настройка логгирования
logger = logging.getLogger()
logger.setLevel(logging.DEBUG)  # Установлено на DEBUG для более подробных сообщений

# Создание обработчика для записи в файл
file_handler = logging.FileHandler('file_processing.log')
file_handler.setLevel(logging.DEBUG)

# Создание обработчика для вывода в консоль
console_handler = logging.StreamHandler()
console_handler.setLevel(logging.INFO)  # Консоль ограничена INFO, чтобы избежать излишней детализации

# Форматтер для сообщений
formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
file_handler.setFormatter(formatter)
console_handler.setFormatter(formatter)

# Добавление обработчиков к логгеру
logger.addHandler(file_handler)
logger.addHandler(console_handler)

# Переменные поступают через main
sourceDirectory = sys.argv[1]
target_path = sys.argv[2]
oldversion_1 = sys.argv[3]
newversion_1 = sys.argv[4]
platformversion_1 = sys.argv[5]
slkversion = sys.argv[6]
sbor = sys.argv[7]


date_year = datetime.now().year
date_today = datetime.now().strftime("%d.%m.%Y")
htm_source_folder = os.path.join(target_path, "Комплект поставки обновления", "Updsetup")
htm_target_folder = os.path.join(target_path, "Комплект первичных материалов", "Update")

# Формирование дополнительных переменных для версии
def create_version_variables(version):
    parts = version.split('.')
    version_2 = '.'.join(parts[:3])
    version_3 = '_'.join(parts)
    return version_2, version_3

oldversion_2, oldversion_3 = create_version_variables(oldversion_1)
newversion_2, newversion_3 = create_version_variables(newversion_1)

# Словарь замен
replacements = {
    "oldversion_1": oldversion_1,
    "newversion_1": newversion_1,
    "platformversion_1": platformversion_1,
    "slkversion": slkversion,
    "date_year": str(date_year),
    "date_today": date_today,
    "oldversion_2": oldversion_2,
    "newversion_2": newversion_2,
    "oldversion_3": oldversion_3,
    "newversion_3": newversion_3,
}

# Универсальная функция для копирования и замены содержимого файлов
def copy_and_process_files(src, dst):
    logger.debug(f"Начало копирования файлов из {src} в {dst}")
    for root, dirs, files in os.walk(src):
        # Создание структуры папок в целевой директории
        dst_root = root.replace(src, dst, 1)
        if not os.path.exists(dst_root):
            os.makedirs(dst_root)
            logger.debug(f"Создана директория: {dst_root}")

        for file in files:
            src_file = os.path.join(root, file)
            dst_file = os.path.join(dst_root, file)

            # Переименование файла на основе замен
            new_file_name = file
            for key, value in replacements.items():
                new_file_name = re.sub(re.escape(key), value, new_file_name)

            dst_file = os.path.join(dst_root, new_file_name)

            try:
                if os.path.exists(dst_file):
                    os.remove(dst_file)  # Удаление существующего файла перед копированием
                    logger.warning(f"Старый файл удален перед копированием: {dst_file}")
                shutil.copy2(src_file, dst_file)  # Копирование файла с сохранением метаданных
                logger.info(f'Файл успешно скопирован и переименован: {src_file} -> {dst_file}')

                # Замена содержимого файла
                if dst_file.endswith('.docx'):
                    replace_text_in_docx_file(dst_file, replacements)
                else:
                    replace_in_text_file(dst_file)

            except Exception as e:
                logger.error(f'Ошибка при копировании и обработке файла {src_file}: {e}')

# Функция для обработки содержимого текстовых файлов
def replace_in_text_file(file_path):
    logger.debug(f"Начало замены текста в файле: {file_path}")
    try:
        with open(file_path, 'r+', encoding='utf-8') as file:
            content = file.read()
            for key, value in replacements.items():
                content = content.replace(key, value)
            file.seek(0)
            file.write(content)
            file.truncate()
            logger.info(f'Замены выполнены в файле: {file_path}')
    except UnicodeDecodeError:
        try:
            with open(file_path, 'r+', encoding='windows-1251') as file:
                content = file.read()
                for key, value in replacements.items():
                    content = content.replace(key, value)
                file.seek(0)
                file.write(content)
                file.truncate()
                logger.info(f'Замены выполнены в файле с кодировкой windows-1251: {file_path}')
        except Exception as e:
            logger.error(f'Ошибка при обработке файла {file_path} с кодировкой windows-1251: {e}')
    except Exception as e:
        logger.error(f'Ошибка при обработке файла {file_path}: {e}')

# Функция для обработки содержимого .docx файлов
def replace_text_in_docx_file(file_path, replacements):
    logger.debug(f"Начало замены текста в .docx файле: {file_path}")
    try:
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            for search_text, replace_text in replacements.items():
                if search_text in paragraph.text:
                    paragraph.text = paragraph.text.replace(search_text, replace_text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for search_text, replace_text in replacements.items():
                            if search_text in paragraph.text:
                                paragraph.text = paragraph.text.replace(search_text, replace_text)
        
        doc.save(file_path)
        logger.info(f'Замены выполнены в файле .docx: {file_path}')
    except Exception as e:
        logger.error(f'Ошибка при обработке файла .docx {file_path}: {e}')

# Функция для поиска и копирования .htm файла
def copy_htm_file(src_folder, dst_folder):
    logger.debug(f"Поиск и копирование .htm файла из {src_folder} в {dst_folder}")
    for root, dirs, files in os.walk(src_folder):
        for file in files:
            if file.lower().endswith('.htm'):
                src_file = os.path.join(root, file)
                dst_file = os.path.join(dst_folder, file)
                try:
                    shutil.copy2(src_file, dst_file)
                    logger.info(f'Файл .htm успешно скопирован: {src_file} -> {dst_file}')
                    return
                except Exception as e:
                    logger.error(f'Ошибка при копировании файла .htm {src_file}: {e}')
    logger.warning('Файл .htm не найден')

# Функция для выполнения переноса файлов
def additional_actions(target_path):
    release_path = rf"{sbor}"
    logger.debug(f"Начало дополнительных действий в {release_path}")

    # Удаление папок "DemoDB", "Update", "ExtFiles" из каталога
    folders_to_delete = ["DemoDB", "Update", "ExtFiles"]
    for folder in folders_to_delete:
        folder_path = os.path.join(release_path, folder)
        if os.path.exists(folder_path):
            shutil.rmtree(folder_path)
            logger.info(f"Папка удалена: {folder_path}")

    # Копирование папок "DemoDB", "Update", "ExtFiles" из "target_path\Комплект первичных материалов" в папку релиза
    folders_to_copy = ["DemoDB", "Update", "ExtFiles"]
    for folder in folders_to_copy:
        src_folder_path = os.path.join(target_path, "Комплект первичных материалов", folder)
        dest_folder_path = os.path.join(release_path, folder)
        if os.path.exists(src_folder_path):
            shutil.copytree(src_folder_path, dest_folder_path)
            logger.info(f"Папка {folder} успешно перенесена")

            # Дополнительное действие для папки "Update"
            if folder == "Update":
                readme_src_path = os.path.join(target_path, r"Комплект поставки обновления\Updsetup\ReadMe.txt")
                readme_dest_path = os.path.join(dest_folder_path, "ReadMe.txt")
                if os.path.exists(readme_src_path):
                    shutil.copyfile(readme_src_path, readme_dest_path)
                    logger.info(f"Файл ReadMe.txt успешно перенесен в папку {folder}")

    # Копирование папки "Protection" из "target_path\Дистрибутив\Configs" в папку релиза
    protection_src_path = os.path.join(target_path, "Дистрибутив", "Configs", "Protection")
    protection_dest_path = os.path.join(release_path,"Update", "Protection")
    if os.path.exists(protection_src_path):
        shutil.copytree(protection_src_path, protection_dest_path)
        logger.info("Папка Protection успешно перенесена")

# Основной блок выполнения
if __name__ == '__main__':
    try:
        copy_and_process_files(sourceDirectory, target_path)
        copy_htm_file(htm_source_folder, htm_target_folder)
        additional_actions(target_path)
        logger.info("Процесс завершен успешно")
    except Exception as e:
        logger.error(f'Ошибка при выполнении основного блока: {e}')
